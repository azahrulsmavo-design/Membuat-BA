from docx import Document
from docx.shared import Mm, Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import requests
from io import BytesIO
import os
from pdf_generator import fetch_drive_image

def fetch_image(path_or_url):
    """
    Fetches an image from a local path or URL and returns a BytesIO object.
    Returns None if the image cannot be loaded.
    """
    if not path_or_url:
        return None

    try:
        # Check if it's a URL
        if path_or_url.startswith('http'):
            # Use the robust Drive fetcher from pdf_generator if it's a drive link or general url
            # The fetch_drive_image function handles generic URLs too if they aren't drive (conceptually, looking at it, it explicitly checks drive patterns, but let's verify)
            # Actually pdf_generator.fetch_drive_image returns None if not drive.
            # So we should check.
            
            if 'drive.google.com' in path_or_url:
                return fetch_drive_image(path_or_url)
            
            # Disable SSL verify for internal consistency with pdf_generator
            response = requests.get(path_or_url, verify=False, timeout=10)
            if response.status_code == 200:
                return BytesIO(response.content)
            else:
                return None
        
        # Check if it's a base64 string (data URI)
        elif path_or_url.startswith('data:image'):
            # Base64 handling
            from base64 import b64decode
            header, encoded = path_or_url.split(",", 1)
            data = b64decode(encoded)
            return BytesIO(data)

        # Local file path
        elif os.path.exists(path_or_url):
             with open(path_or_url, 'rb') as f:
                return BytesIO(f.read())
        
        return None

    except Exception as e:
        print(f"Error fetching image {path_or_url}: {e}")
        return None

def set_cell_margins(cell, top=0, start=0, bottom=0, end=0):
    """
    Sets cell margins for a specific table cell.
    """
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    
    for name, value in [('top', top), ('left', start), ('bottom', bottom), ('right', end)]:
        node = OxmlElement(f'w:{name}')
        node.set(qn('w:w'), str(int(value * 20))) # value in points * 20 = twips? No, input is roughly pixels or small units? 
        # Actually docx uses distinct units. Let's rely on standard paragraph spacing inside calls instead if this is complex.
        # But this function is useful for strict layout. 
        # width in w:w is in twentieths of a point (dxa).
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    
    tcPr.append(tcMar)

def create_multiset_docx(data, output_path):
    doc = Document()
    
    # Set Narrow Margins (1.27 cm)
    # Set Narrow Margins (1.0 cm) to match PDF 190mm content width logic
    section = doc.sections[0]
    section.page_height = Mm(297)
    section.page_width = Mm(210)
    section.left_margin = Mm(10)
    section.right_margin = Mm(10)
    section.top_margin = Mm(10)
    section.bottom_margin = Mm(10)

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Calibri'
    font.size = Pt(11)

    units = data.get('units', [])
    layout_config = data.get('layout', {}) 

    for i, unit in enumerate(units):
        nopol = unit.get('nopol', 'UNKNOWN')
        bu = unit.get('bu', 'UNKNOWN')
        lokasi = unit.get('lokasi', 'UNKNOWN')
        images = unit.get('images', {})

        # --- PAGE 1: Check Fisik Dokumen ---
        
        # Header - Matches PDF: "BU : {bu} - {location} - {nopol}"
        header_title = f"BU : {bu} - {lokasi} - {nopol}".upper()
        
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(header_title)
        run.bold = True
        run.font.size = Pt(11) # PDF uses 11
        
        # Space after header
        doc.add_paragraph() 

        # doc.add_heading('1. DOKUMEN KENDARAAN', level=2) # Removed to match PDF clean look where visual boxes dominate, or keep if user wants headings?
        # PDF code has: pdf.cell(..., "FOTO STNK...", ...) inside the box logic.
        # It DOES NOT have a main heading "1. DOKUMEN KENDARAAN".
        # It just starts.
        
        # But we need some separation. The PDF has "BU : ..." then "FOTO STNK (SURAT TANDA NOMOR KENDARAAN) :"
        
        # Table for Documents (STNK, Pajak, KIR)
        # Layout: 
        # Row 1: STNK (Full Width)
        # Row 2: Pajak (Full Width)
        # Row 3: KIR Kertas (Left), KIR Kartu (Right)
        
        table = doc.add_table(rows=3, cols=2)
        table.style = 'Table Grid'
        table.autofit = False
        
        # Set column widths (Total 190mm)
        # Col 1: 95mm, Col 2: 95mm (roughly, gap included in cell padding usually)
        for row in table.rows:
            for cell in row.cells:
                cell.width = Mm(95)

        # Helper to fill cell
        def fill_cell(cell, title, img_key, max_width_mm=85):
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run(title + "\n")
            run.bold = True
            
            img_data = images.get(img_key)
            if img_data:
                # Handle both dict with dataUrl or direct string
                path = img_data.get('dataUrl') if isinstance(img_data, dict) else img_data
                
                img_stream = fetch_image(path)
                if img_stream:
                    try:
                        # Add image, constraining width
                        run.add_picture(img_stream, width=Mm(max_width_mm))
                    except Exception as e:
                        print(f"Error adding picture {img_key}: {e}")
                        run.add_text(f"[Error loading image]")
                else:
                    run.add_text("[Gambar Tidak Ditemukan]")
            else:
                run.add_text("[Tidak Ada Gambar]")

        # Row 1: STNK (Merge cells)
        cell_stnk = table.cell(0, 0)
        cell_stnk.merge(table.cell(0, 1))
        # Height constraint approx 85mm
        fill_cell(cell_stnk, "FOTO STNK (SURAT TANDA NOMOR KENDARAAN) :", 'stnk', max_width_mm=160)
        
        # Row 2: Pajak (Merge cells)
        cell_pajak = table.cell(1, 0)
        cell_pajak.merge(table.cell(1, 1))
        fill_cell(cell_pajak, "FOTO LEMBAR PAJAK :", 'tax', max_width_mm=160)
        
        # Row 3: KIR (Split)
        # Height approx 65mm
        fill_cell(table.cell(2, 0), "FOTO LEMBAR KIR :", 'kir', max_width_mm=85)
        fill_cell(table.cell(2, 1), "FOTO KARTU KIR :", 'kir_card', max_width_mm=85)

        doc.add_page_break()

        # --- PAGE 2: Foto Fisik Kendaraan ---
        
        # Header (Repeated for context)
        # pdf_generator page 2 has header too
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        run = p.add_run(header_title)
        run.bold = True
        run.font.size = Pt(11)
        
        doc.add_paragraph() # Spacer
        
        # doc.add_heading('2. KONDISI FISIK', level=2) # Removed to match PDF

        # Table for Physical Photos
        # Row 1: Depan, Belakang
        # Row 2: Kanan, Kiri
        
        table_phys = doc.add_table(rows=2, cols=2)
        table_phys.style = 'Table Grid'
        
        fill_cell(table_phys.cell(0, 0), "TAMPAK DEPAN", 'front', max_width_mm=85)
        fill_cell(table_phys.cell(0, 1), "TAMPAK BELAKANG", 'back', max_width_mm=85)
        fill_cell(table_phys.cell(1, 0), "TAMPAK SAMPING KANAN", 'right', max_width_mm=85)
        fill_cell(table_phys.cell(1, 1), "TAMPAK SAMPING KIRI", 'left', max_width_mm=85)

        doc.add_page_break()

    # --- PAGE 3: Summary (Rekapitulasi) ---
    doc.add_heading('REKAPITULASI ASET KENDARAAN', level=1)
    
    summary_table = doc.add_table(rows=1, cols=4)
    summary_table.style = 'Table Grid'
    
    # Header Row
    hdr_cells = summary_table.rows[0].cells
    hdr_cells[0].text = 'No'
    hdr_cells[1].text = 'Nomor Polisi'
    hdr_cells[2].text = 'Status Dokumen'
    hdr_cells[3].text = 'Kekurangan'
    
    for i, unit in enumerate(units):
        row_cells = summary_table.add_row().cells
        row_cells[0].text = str(i + 1)
        row_cells[1].text = unit.get('nopol', 'UNKNOWN')
        
        # Logic to check completeness
        unit_images = unit.get('images', {})
        required_keys = ['stnk', 'tax', 'front', 'back', 'right', 'left'] # Basic Requirement
        missing = []
        
        for key in required_keys:
            val = unit_images.get(key)
            # Check for null, empty string, or empty dict
            if not val or (isinstance(val, dict) and not val.get('dataUrl')):
                missing.append(key.upper())
        
        if not missing:
            row_cells[2].text = "LENGKAP"
            row_cells[3].text = "-"
        else:
            row_cells[2].text = "TIDAK LENGKAP"
            row_cells[3].text = ", ".join(missing)

    doc.save(output_path)
    return output_path
